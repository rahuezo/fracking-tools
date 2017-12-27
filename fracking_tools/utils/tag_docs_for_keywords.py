from files import ExtensionHandler, get_text
from docx import Document

import csv
import re


class KeywordTagger:
    def __init__(self, f, keywords, output_path):
        self.file_to_tag = f
        # self.keywords_file = keywords_file
        self.keywords = keywords #self.get_keywords()
        self.output_path = output_path
        self.content = ExtensionHandler(self.file_to_tag).get_content()

    def keyword_mapper(self, content, tags):
        for i in xrange(len(self.keywords)):
            kw = self.keywords[i]
            if re.findall(r'\b{0}\b'.format(kw), content.lower()):
                content = get_text(content.lower().replace(kw, tags[i]))
        return content

    # def get_keywords(self):
    #     with open(self.keywords_file, 'r') as kwds_file:
    #         return map(get_text, kwds_file.read().split(','))

    def find_keywords(self, content, tags):
        return filter(lambda tag: re.findall(r'{0}'.format(tag), content), tags)

    def generate_output(self):
        tags = [' <key{0}> '.format(i) for i in xrange(len(self.keywords))]
        content = self.keyword_mapper(self.content, tags)
        results = self.find_keywords(content, tags)

        if results:
            doc = Document()
            doc.add_heading(', '.join(map(lambda x: x.title(),
                                          map(lambda x: self.keywords[tags.index(x)], results))), 0)
            split_content = content.split()
            paragraph = doc.add_paragraph('')

            for word in split_content:
                if word.lower() in map(lambda x: x.strip(), tags):
                    paragraph.add_run(' {0}'.format(self.keywords[tags.index(' {0} '.format(word.lower()))].upper())).bold = True
                else:
                    paragraph.add_run(' {0}'.format(word))

            doc.save(self.output_path)
            return self.output_path
        else:
            return None



# kt = KeywordTagger('content.docx', 'kwds.txt', 'sample.docx')
#
# kt.generate_output()






